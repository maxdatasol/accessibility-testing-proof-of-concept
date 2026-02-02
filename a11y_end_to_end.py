#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
End-to-end A11Y pipeline (consolidated document + front-matter)

- Input: array of URLs (via --urls / --urls-json / --file)
- Run axe-core via Selenium → a11y_violations.csv
- Run Lighthouse via npx → lighthouse_summary.csv + lighthouse_failing_audits.csv (optional)
- Build ONE consolidated Word document from IssueTemplate.docx (one section per violation rule),
  prefixed with a front-matter Word document (cover/summary/urls/conclusion) if provided.

Requirements:
    pip install selenium axe-selenium-python webdriver-manager python-docx
    Node.js + npx for Lighthouse (optional)
To execute: C:\path\to\venv\Scripts\python.exe a11y_end_to_end.py --urls-json "[\"http://max-data.com\"]" --axe --headless
"""

import argparse, csv, json, os, re, subprocess, sys, time, copy
from collections import Counter, defaultdict
from datetime import datetime
from shutil import which
from typing import Dict, Iterable, List, Optional, Tuple

# ---- axe / Selenium ----
from selenium import webdriver
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from axe_selenium_python import Axe

# ---- Word ----
from docx import Document
from docx.enum.text import WD_BREAK

IMPACT_ORDER = ["minor", "moderate", "serious", "critical"]

# -------------------
# Utilities
# -------------------

def build_issue_fields(violation_id: str, records: List[Dict[str, str]], lh_by_url: Dict[str, Dict[str, str]]) -> Dict[str, str]:
    # WCAG info
    sc_set = set()
    level_set = set()
    for r in records:
        tags_str = r.get("tags", "")
        sc, level = wcag_from_tags(tags_str)
        if sc: sc_set.add(sc)
        if level: level_set.add(level)
    success_criteria = ", ".join(sorted(sc_set)) or "n/a"
    success_level = ", ".join(sorted(level_set)) or "n/a"

    # Description fields
    issue_descriptions = set()
    for r in records:
        desc = r.get("description", "")
        help_text = r.get("help", "")
        combined = " — ".join(filter(None, [desc, help_text]))
        if combined:
            issue_descriptions.add(combined)
    issue_description = "\n\n".join(sorted(issue_descriptions)) or "n/a"
    success_criteria_description = issue_description  # same as description

    # Long description
    long_descriptions = set(r.get("failure_summary", "") for r in records if r.get("failure_summary"))
    issue_long_description = "\n\n".join(sorted(long_descriptions)) or "n/a"

    # Pages affected
    pages = set(r.get("page_url") for r in records if r.get("page_url"))
    pages_affected = "\n".join(sorted(pages)) or "n/a"

    # HTML code
    html_snippets = set(r.get("html", "") for r in records if r.get("html"))
    issue_code = "\n\n".join(sorted(html_snippets)) or "n/a"

    # Implications / impact
    impacts = set(r.get("impact", "").capitalize() for r in records if r.get("impact"))
    issue_implications = "Impact: " + ", ".join(sorted(impacts)) if impacts else "n/a"

    # Solution / help URL
    help_urls = set(r.get("help_url", "") for r in records if r.get("help_url"))
    issue_solution = "\n".join(sorted(help_urls)) or "n/a"

    # Related violations / Lighthouse notes
    related = []
    for r in records:
        url = r.get("page_url")
        lh = lh_by_url.get(url, {})
        if lh:
            related.append(f"{url} → Accessibility score: {lh.get('score_accessibility', 'n/a')}")
    related_violations = "\n".join(related) if related else "n/a"

    return {
        "SuccessCriteria": success_criteria,
        "SuccessCriterionLevel": success_level,
        "SuccessCriteriaDescription": success_criteria_description,
        "IssueDescription": issue_description,
        "IssueLongDescription": issue_long_description,
        "PagesAffected": pages_affected,
        "IssueCode": issue_code,
        "IssueImplications": issue_implications,
        "IssueSolution": issue_solution,
        "RelatedViolations": related_violations,
    }


def parse_urls(urls_arg: Optional[str], urls_json: Optional[str], file_arg: Optional[str]) -> List[str]:
    urls: List[str] = []
    if urls_arg:
        urls.extend([u.strip() for u in urls_arg.split(",") if u.strip()])
    if urls_json:
        try:
            arr = json.loads(urls_json)
            if not isinstance(arr, list):
                raise ValueError("--urls-json must be a JSON array")
            urls.extend([str(u).strip() for u in arr if str(u).strip()])
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON in --urls-json: {e}")
    if file_arg:
        with open(file_arg, "r", encoding="utf-8") as f:
            for line in f:
                s = line.strip()
                if s and not s.startswith("#"):
                    urls.append(s)
    # De-dupe preserve order
    seen = set()
    out: List[str] = []
    for u in urls:
        if u not in seen:
            out.append(u)
            seen.add(u)
    return out

def write_csv(path: str, rows: Iterable[Dict[str, str]], fieldnames: List[str]):
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        w.writeheader()
        for r in rows:
            w.writerow(r)

def impact_at_least(impact: Optional[str], min_impact: Optional[str]) -> bool:
    if not min_impact:
        return True
    if impact is None:
        return False
    try:
        return IMPACT_ORDER.index(impact) >= IMPACT_ORDER.index(min_impact)
    except ValueError:
        return False

# -------------------
# axe via Selenium
# -------------------

def get_driver(headless: bool = True) -> webdriver.Chrome:
    options = ChromeOptions()
    if headless:
        options.add_argument("--headless=new")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--window-size=1920,1080")
    service = ChromeService(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=options)

def run_axe_for_url(driver: webdriver.Chrome, url: str, wait_seconds: float, tags: Optional[List[str]], include_iframes: bool) -> Dict:
    driver.get(url)
    if wait_seconds > 0:
        time.sleep(wait_seconds)
    axe = Axe(driver)
    axe.inject()
    options = {
        "resultTypes": ["violations"],
        **({"runOnly": {"type": "tag", "values": tags}} if tags else {}),
    }
    if include_iframes:
        options["iframes"] = True
    return axe.run(options=options)

def axe_to_rows(violations: List[Dict], page_url: str, page_title: str, min_impact: Optional[str]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for v in violations:
        v_impact = v.get("impact")
        if not impact_at_least(v_impact, min_impact):
            continue
        base = {
            "page_url": page_url,
            "page_title": page_title or "",
            "violation_id": v.get("id", ""),
            "impact": v_impact or "",
            "help": v.get("help", ""),
            "help_url": v.get("helpUrl", ""),
            "description": v.get("description", ""),
            "tags": ";".join(v.get("tags", []) or []),
        }
        nodes = v.get("nodes", []) or []
        if not nodes:
            rows.append({**base, "target": "", "failure_summary": "", "html": ""})
            continue
        for node in nodes:
            rows.append({
                **base,
                "target": " | ".join(node.get("target", []) or []),
                "failure_summary": node.get("failureSummary", "") or "",
                "html": (node.get("html", "") or "").replace("\n", " ").strip(),
            })
    return rows

# -------------------
# Lighthouse via npx
# -------------------

def resolve_cmd(cmd: str) -> Optional[str]:
    p = which(cmd)
    if p: return p
    if os.name == "nt":
        p = which(cmd + ".cmd")
        if p: return p
    return None

def build_lighthouse_cmd(url: str, categories: Optional[List[str]], preset: str, chrome_flags: Optional[List[str]], quiet: bool = True) -> List[str]:
    npx = resolve_cmd("npx")
    if not npx:
        raise FileNotFoundError("npx not found on PATH. Install Node.js or add npx to PATH.")
    cmd = [npx, "lighthouse", url, "--output=json", "--output-path=stdout", f"--preset={preset}"]
    if categories:
        cmd += ["--only-categories=" + ",".join(categories)]
    flags = chrome_flags or ["--headless=new", "--no-sandbox", "--disable-dev-shm-usage", "--window-size=1920,1080"]
    cmd.append("--chrome-flags=" + " ".join(flags))
    if quiet:
        cmd.append("--quiet")
    return cmd

def run_lighthouse(url: str, categories: Optional[List[str]], preset: str, chrome_flags: Optional[List[str]]) -> Optional[Dict]:
    try:
        cmd = build_lighthouse_cmd(url, categories, preset, chrome_flags, quiet=True)
    except FileNotFoundError:
        return None
    try:
        p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True)
        return json.loads(p.stdout)
    except subprocess.CalledProcessError as e:
        try:
            return json.loads(e.stdout or "")
        except Exception:
            sys.stderr.write(f"[LH] Failed for {url}: {e.stderr}\n")
            return None
    except json.JSONDecodeError as e:
        sys.stderr.write(f"[LH] JSON parse failed for {url}: {e}\n")
        return None

def lh_summary_row(url: str, lh: Dict) -> Dict[str, str]:
    cats = lh.get("categories", {}) or {}
    def score(cat):
        try:
            return int(round((cats.get(cat, {}).get("score") or 0) * 100))
        except Exception:
            return ""
    env = lh.get("environment", {}) or {}
    ua = lh.get("userAgent", "") or env.get("hostUserAgent", "")
    return {
        "page_url": url,
        "final_url": lh.get("finalUrl", ""),
        "fetch_time": lh.get("fetchTime", ""),
        "user_agent": ua,
        "device": lh.get("configSettings", {}).get("formFactor", ""),
        "score_performance": score("performance"),
        "score_accessibility": score("accessibility"),
        "score_best_practices": score("best-practices"),
        "score_seo": score("seo"),
        "score_pwa": score("pwa"),
    }

# -------------------
# Word: per-violation section generation
# -------------------

def replace_all_text(doc: Document, repl: Dict[str, str]):
    """
    Replace placeholders in doc, even if Word splits them across multiple runs.
    """
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, repl)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, repl)

def _replace_in_paragraph(paragraph, repl):
    """
    Replace placeholders in a single paragraph by merging all runs temporarily.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    replaced = full_text
    for key, val in repl.items():
        replaced = replaced.replace(key, val)
    if replaced != full_text:
        # Remove old runs
        for run in paragraph.runs:
            run.text = ""
        # Add single run with replaced text
        paragraph.runs[0].text = replaced


def wcag_from_tags(tags_str: str) -> Tuple[str, str]:
    tags = [(t or '').strip().lower() for t in (tags_str or '').split(';') if t.strip()]
    level_map = {'wcag2a':'WCAG 2.0 A','wcag2aa':'WCAG 2.0 AA','wcag21a':'WCAG 2.1 A','wcag21aa':'WCAG 2.1 AA','wcag22a':'WCAG 2.2 A','wcag22aa':'WCAG 2.2 AA'}
    level, sc = '', ''
    for key in ['wcag22aa','wcag22a','wcag21aa','wcag21a','wcag2aa','wcag2a']:
        if key in tags:
            level = level_map[key]
            break
    for t in tags:
        m = re.fullmatch(r'wcag(\d{3,4})', t)
        if m:
            digits = m.group(1)
            sc = f"{digits[0]}.{digits[1]}.{digits[2:]}"
            break
    return sc, level

# (keep build_issue_fields as-is)
# -------------------
# Front matter helpers
# -------------------

def which_exe(candidates: List[str]) -> Optional[str]:
    for c in candidates:
        path = which(c)
        if path: return path
        if os.name == 'nt':
            path = which(c + '.exe') or which(c + '.cmd')
            if path: return path
    return None

def try_convert_doc_to_docx(doc_path: str) -> str:
    lp = doc_path.lower()
    if not lp.endswith('.doc') or lp.endswith('.docx'):
        return doc_path
    soffice = which_exe(['soffice','libreoffice'])
    if not soffice:
        raise FileNotFoundError('Front-matter is a .doc file. Install LibreOffice (soffice) or provide .docx')
    out_dir = os.path.dirname(os.path.abspath(doc_path)) or '.'
    cmd = [soffice, '--headless', '--convert-to','docx','--outdir', out_dir, doc_path]
    res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if res.returncode != 0:
        raise RuntimeError(f'LibreOffice conversion failed: {res.stderr}')
    new_path = os.path.splitext(doc_path)[0] + '.docx'
    if not os.path.exists(new_path):
        raise RuntimeError(f'Converted file not found: {new_path}')
    return new_path

def add_heading(doc: Document, text: str, level: int = 1):
    try:
        doc.add_heading(text, level=level)
    except Exception:
        p = doc.add_paragraph()
        p.add_run(text).bold = True

def apply_front_matter(front_path: str, urls: List[str], axe_rows: List[Dict[str,str]], lh_by_url: Dict[str,Dict[str,str]], client_name: Optional[str], report_date: Optional[str], conclusion_text: Optional[str]) -> Document:
    front_docx = try_convert_doc_to_docx(front_path)
    base = Document(front_docx)

    repl = {}
    if client_name: repl['«ClientName»'] = client_name
    repl['«ReportDate»'] = report_date or datetime.now().strftime('%Y-%m-%d')
    if repl: replace_all_text(base, repl)

    # Summary
    summary_lines = [
        f"URLs tested: {len(urls)}",
        f"Total axe violation instances: {len(axe_rows)}",
        "Impact distribution: " + ", ".join(f"{k}={v}" for k,v in Counter((r.get('impact') or '').lower() for r in axe_rows).items()) if axe_rows else "n/a"
    ]
    add_heading(base, "Executive Summary", level=1)
    for line in summary_lines:
        base.add_paragraph(line)

    # URLs
    add_heading(base, "URLs Tested", level=1)
    for u in urls:
        base.add_paragraph(f"• {u}")

    # Optional conclusion
    if conclusion_text:
        add_heading(base, "Conclusion", level=1)
        base.add_paragraph(conclusion_text)

    return base

# -------------------
# Main
# -------------------

def main():
    parser = argparse.ArgumentParser(description="End-to-end accessibility pipeline")
    parser.add_argument("--urls", help="Comma-separated URLs")
    parser.add_argument("--urls-json", help="JSON array of URLs")
    parser.add_argument("--file", help="File with URLs (one per line)")
    parser.add_argument("--axe", action="store_true")
    parser.add_argument("--lighthouse", action="store_true")
    parser.add_argument("--headless", action="store_true")
    parser.add_argument("--template", help="Issue template DOCX")
    parser.add_argument("--front-matter", help="Front matter DOCX")
    parser.add_argument("--client-name")
    parser.add_argument("--report-date")
    parser.add_argument("--single-doc-out", help="Output Word document")
    parser.add_argument("--wait", type=float, default=0)
    parser.add_argument("--min-impact")
    parser.add_argument("--tags", nargs="*")
    parser.add_argument("--include-iframes", action="store_true")
    args = parser.parse_args()

    urls = parse_urls(args.urls, args.urls_json, args.file)
    if not urls:
        print("No URLs provided — exiting.")
        return

    print(f"URLs to test: {len(urls)}")
    axe_rows = []
    lh_by_url = {}
    driver = None

    try:
        if args.axe:
            driver = get_driver(headless=args.headless)
            for url in urls:
                print(f"[AXE] Scanning {url}")
                results = run_axe_for_url(driver, url, wait_seconds=args.wait, tags=args.tags, include_iframes=args.include_iframes)
                violations = results.get("violations", [])
                axe_rows.extend(axe_to_rows(violations, page_url=url, page_title=driver.title, min_impact=args.min_impact))

        if args.lighthouse:
            for url in urls:
                print(f"[LH] Auditing {url}")
                lh = run_lighthouse(url, ["accessibility"], "desktop", None)
                if lh:
                    lh_by_url[url] = lh_summary_row(url, lh)

        if args.single_doc_out and args.template:
            print("[DOC] Building Word report")
            doc = Document()
            if args.front_matter:
                doc = apply_front_matter(args.front_matter, urls, axe_rows, lh_by_url, args.client_name, args.report_date, None)

            by_violation = defaultdict(list)
            
            for r in axe_rows:
                by_violation[r["violation_id"]].append(r)

            for vid, records in by_violation.items():
                doc.add_page_break()
                issue_doc = Document(args.template)
                fields = build_issue_fields(vid, records, lh_by_url)
                replace_all_text(issue_doc, fields)
                for element in issue_doc.element.body:
                    doc.element.body.append(copy.deepcopy(element))

            doc.save(args.single_doc_out)
            print(f"[DONE] Report written to {args.single_doc_out}")

    finally:
        if driver:
            driver.quit()

if __name__ == "__main__":
    main()
